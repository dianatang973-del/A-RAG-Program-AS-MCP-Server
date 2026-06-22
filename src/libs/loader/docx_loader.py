"""Word Document (DOCX) Loader implementation.

This module implements DOCX file parsing with conversion to Markdown.

Features:
- DOCX to Markdown conversion via mammoth
- Image extraction from document
- Title extraction from document properties or first heading
- Paragraph and formatting preservation
"""

from __future__ import annotations

import hashlib
import logging
from pathlib import Path
from typing import Any, Dict, List, Optional
import io

try:
    import mammoth
    MAMMOTH_AVAILABLE = True
except ImportError:
    MAMMOTH_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import _Cell, Table
    from docx.text.paragraph import Paragraph
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

from PIL import Image

from src.core.types import Document
from src.libs.loader.base_loader import BaseLoader

logger = logging.getLogger(__name__)


class DocxLoader(BaseLoader):
    """Word Document (DOCX) Loader using mammoth for Markdown conversion.

    This loader:
    1. Converts DOCX to Markdown using mammoth
    2. Extracts images from document
    3. Saves images to storage directory
    4. Inserts image placeholders in the format [IMAGE: {image_id}]
    5. Records image metadata in Document.metadata.images

    Configuration:
        extract_images: Enable/disable image extraction (default: True)
        image_storage_dir: Base directory for image storage (default: data/images)

    Dependencies:
        - mammoth: For DOCX to Markdown conversion
        - python-docx: For image extraction and metadata
        - Pillow: For image processing

    Note: mammoth provides better Markdown conversion than python-docx alone.
    """

    def __init__(
        self,
        extract_images: bool = True,
        image_storage_dir: str | Path = "data/images"
    ):
        """Initialize DOCX Loader.

        Args:
            extract_images: Whether to extract images from DOCX.
            image_storage_dir: Base directory for storing extracted images.

        Raises:
            ImportError: If required dependencies are not installed.
        """
        if not MAMMOTH_AVAILABLE:
            raise ImportError(
                "mammoth is required for DocxLoader. "
                "Install with: pip install mammoth"
            )

        if not PYTHON_DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for DocxLoader. "
                "Install with: pip install python-docx"
            )

        self.extract_images = extract_images
        self.image_storage_dir = Path(image_storage_dir)

    def load(self, file_path: str | Path) -> Document:
        """Load and parse a DOCX file.

        Args:
            file_path: Path to the DOCX file.

        Returns:
            Document with Markdown text and metadata.

        Raises:
            FileNotFoundError: If the DOCX file doesn't exist.
            ValueError: If the file is not a valid DOCX file.
            RuntimeError: If parsing fails critically.
        """
        # Validate file
        path = self._validate_file(file_path)
        if path.suffix.lower() not in ['.docx']:
            raise ValueError(f"File is not a DOCX file: {path}")

        # Compute document hash for unique ID
        doc_hash = self._compute_file_hash(path)
        doc_id = f"doc_{doc_hash[:16]}"

        # Convert DOCX to Markdown using mammoth
        # Note: mammoth automatically converts tables to Markdown table format
        # No additional configuration needed for table preservation
        try:
            with open(path, 'rb') as docx_file:
                result = mammoth.convert_to_markdown(docx_file)
                markdown_text = result.value
                conversion_messages = result.messages

            # Log any conversion warnings (including table conversion issues)
            for message in conversion_messages:
                if message.type == 'warning':
                    logger.warning(f"DOCX conversion warning: {message.message}")

        except Exception as e:
            logger.error(f"Failed to convert DOCX to Markdown {path}: {e}")
            raise RuntimeError(f"DOCX to Markdown conversion failed: {e}") from e

        # Initialize metadata
        metadata: Dict[str, Any] = {
            "source_path": str(path),
            "doc_type": "docx",
            "doc_hash": doc_hash,
        }

        # Extract title and other metadata using python-docx
        try:
            docx_doc = DocxDocument(str(path))

            # Try to get title from document properties
            title = self._extract_title(docx_doc, markdown_text)
            if title:
                metadata["title"] = title

            # Extract paragraph count as a proxy for "pages"
            # (Word doesn't have fixed pages like PDF)
            paragraph_count = len(docx_doc.paragraphs)
            metadata["paragraph_count"] = paragraph_count

        except Exception as e:
            logger.warning(f"Failed to extract metadata from {path}: {e}")

        # Handle image extraction
        if self.extract_images:
            try:
                markdown_text, images_metadata = self._extract_and_process_images(
                    path, markdown_text, doc_hash
                )
                if images_metadata:
                    metadata["images"] = images_metadata
            except Exception as e:
                logger.warning(
                    f"Image extraction failed for {path}, continuing with text-only: {e}"
                )

        return Document(
            id=doc_id,
            text=markdown_text,
            metadata=metadata
        )

    def _compute_file_hash(self, file_path: Path) -> str:
        """Compute SHA256 hash of file content."""
        sha256 = hashlib.sha256()
        with open(file_path, 'rb') as f:
            for chunk in iter(lambda: f.read(8192), b''):
                sha256.update(chunk)
        return sha256.hexdigest()

    def _extract_title(self, docx_doc: DocxDocument, markdown_text: str) -> Optional[str]:
        """Extract title from DOCX document.

        Priority:
        1. Document core properties title
        2. First heading in Markdown
        3. First paragraph

        Args:
            docx_doc: python-docx Document object.
            markdown_text: Converted Markdown text.

        Returns:
            Title string if found, None otherwise.
        """
        # Try document properties
        try:
            if docx_doc.core_properties.title:
                return docx_doc.core_properties.title.strip()
        except Exception:
            pass

        # Try first heading from Markdown
        lines = markdown_text.split('\n')
        for line in lines[:20]:
            line = line.strip()
            if line.startswith('# '):
                return line[2:].strip()

        # Try first paragraph
        if docx_doc.paragraphs:
            first_para = docx_doc.paragraphs[0].text.strip()
            if first_para:
                return first_para[:100]  # Limit length

        return None

    def _extract_and_process_images(
        self,
        docx_path: Path,
        markdown_text: str,
        doc_hash: str
    ) -> tuple[str, List[Dict[str, Any]]]:
        """Extract images from DOCX and insert placeholders.

        Args:
            docx_path: Path to DOCX file.
            markdown_text: Converted Markdown text.
            doc_hash: Document hash for image ID generation.

        Returns:
            Tuple of (modified_text, images_metadata_list)
        """
        if not self.extract_images:
            logger.debug(f"Image extraction disabled for {docx_path}")
            return markdown_text, []

        images_metadata = []
        modified_text = markdown_text

        try:
            # Create image storage directory
            image_dir = self.image_storage_dir / doc_hash
            image_dir.mkdir(parents=True, exist_ok=True)

            # Open document with python-docx
            docx_doc = DocxDocument(str(docx_path))

            # Extract images from document relationships
            img_index = 0
            for rel in docx_doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        # Get image data
                        image_data = rel.target_part.blob

                        # Determine image extension
                        content_type = rel.target_part.content_type
                        ext = self._get_image_extension(content_type)

                        # Generate image ID and filename
                        image_id = self._generate_image_id(doc_hash, img_index + 1)
                        image_filename = f"{image_id}.{ext}"
                        image_path = image_dir / image_filename

                        # Save image
                        with open(image_path, 'wb') as img_file:
                            img_file.write(image_data)

                        # Get image dimensions
                        try:
                            img = Image.open(io.BytesIO(image_data))
                            width, height = img.size
                        except Exception:
                            width, height = 0, 0

                        # Create placeholder
                        placeholder = f"\n[IMAGE: {image_id}]\n"

                        # Append placeholder at the end
                        # (mammoth doesn't preserve exact image positions)
                        text_offset = len(modified_text)
                        modified_text += placeholder

                        # Convert path to be relative to project root or absolute
                        try:
                            relative_path = image_path.relative_to(Path.cwd())
                        except ValueError:
                            relative_path = image_path.absolute()

                        # Record metadata
                        image_metadata = {
                            "id": image_id,
                            "path": str(relative_path),
                            "text_offset": text_offset,
                            "text_length": len(placeholder.strip()),
                            "position": {
                                "width": width,
                                "height": height,
                                "index": img_index
                            }
                        }
                        images_metadata.append(image_metadata)

                        logger.debug(f"Extracted image {image_id} from DOCX")

                        img_index += 1

                    except Exception as e:
                        logger.warning(f"Failed to extract image {img_index}: {e}")
                        continue

            if images_metadata:
                logger.info(f"Extracted {len(images_metadata)} images from {docx_path}")

        except Exception as e:
            logger.warning(f"Image extraction failed for {docx_path}: {e}")
            return markdown_text, []

        return modified_text, images_metadata

    @staticmethod
    def _get_image_extension(content_type: str) -> str:
        """Get file extension from MIME content type.

        Args:
            content_type: MIME type (e.g., 'image/png').

        Returns:
            File extension (e.g., 'png').
        """
        mime_to_ext = {
            'image/png': 'png',
            'image/jpeg': 'jpg',
            'image/jpg': 'jpg',
            'image/gif': 'gif',
            'image/bmp': 'bmp',
            'image/tiff': 'tiff',
            'image/webp': 'webp',
        }
        return mime_to_ext.get(content_type.lower(), 'png')

    @staticmethod
    def _generate_image_id(doc_hash: str, sequence: int) -> str:
        """Generate unique image ID."""
        return f"{doc_hash[:8]}_img_{sequence}"
